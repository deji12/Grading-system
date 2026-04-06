from utils import save_to_csv, CLASSES_ROOT_DIR, load_config
import os
from pathlib import Path
import json
import csv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import time

def _get_class_info(data):
    cls = data.get("class") or {}
    class_name = (cls.get("name") or "").strip()
    class_group = (cls.get("group") or "").strip()
    return class_name, class_group, f"{class_name}{class_group}"


def _get_subject_lookup(config):
    return {s.get("id"): s.get("name") for s in config.get("subjects", [])}


def _write_class_csv(class_name, class_group, filename, data, field_names):
    save_to_csv(
        data=data,
        file_destination=os.path.join(
            CLASSES_ROOT_DIR,
            f"{class_name}/results/{class_name}_{class_group}/{filename}"
        ),
        field_names=field_names,
    )


def _write_overall_csv(class_name, filename, data, field_names):
    save_to_csv(
        data=data,
        file_destination=os.path.join(
            CLASSES_ROOT_DIR,
            f"{class_name}/results/{filename}"
        ),
        field_names=field_names,
    )


def _count_best_subjects_per_student(subject_toppers):
    student_subjects = {}
    for item in subject_toppers:
        name = (item.get("name") or "").strip()
        subject_id = item.get("subject_id")
        if not name or subject_id is None:
            continue
        student_subjects.setdefault(name, set()).add(subject_id)
    return student_subjects


def _write_section_table(doc, section_name, rows):
    """Add a section heading and a table with the given rows."""
    heading = doc.add_heading(section_name, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    col_headers = ["S/N", "Class", "Name", "Subject(s)", "Average", "Position", "Improvement"]
    table = doc.add_table(rows=1 + len(rows), cols=len(col_headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Optional: set column widths
    for col in table.columns:
        col.width = Inches(1.0)
    table.columns[0].width = Inches(0.5)   # S/N
    table.columns[2].width = Inches(1.5)   # Name
    table.columns[3].width = Inches(2.0)   # Subject(s)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(col_headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Data rows
    for idx, row_data in enumerate(rows, start=1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = str(idx)
        row_cells[1].text = row_data.get("Class", "")
        row_cells[2].text = row_data.get("Name", "")
        row_cells[3].text = row_data.get("Subject(s)", "")
        row_cells[4].text = str(row_data.get("Average", ""))
        row_cells[5].text = row_data.get("Position", "")
        row_cells[6].text = row_data.get("Improvement", "")

    doc.add_paragraph()

def generate_word_report(csv_path, output_path, title):
    """Generate a Word report from a summary CSV file."""
    rows = []
    with open(csv_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append(row)

    if not rows:
        print(f"No data in {csv_path}, skipping Word report.")
        return True

    sections = {}
    for row in rows:
        section = row.get("Section", "Other")
        sections.setdefault(section, []).append(row)

    doc = Document()
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section_name, section_rows in sections.items():
        _write_section_table(doc, section_name, section_rows)

    max_retries = 3
    for attempt in range(max_retries):
        try:
            doc.save(output_path)
            print(f"Word report saved: {output_path}")
            return True
        except PermissionError as e:
            if attempt < max_retries - 1:
                print(f"Permission denied, retrying in 1 second... (attempt {attempt+1}/{max_retries})")
                time.sleep(1)
            else:
                error_msg = (f"Cannot save '{output_path}'. The file is probably open in another program. "
                             f"Please close it and try again.\n\nOriginal error: {str(e)}")
                raise PermissionError(error_msg)
    return False

def generate_all_word_reports():
    """Generate Word reports for JS1-3 and SS1-3 from the summary CSVs."""
    js_csv = os.path.join(CLASSES_ROOT_DIR, "JS1-3.csv")
    ss_csv = os.path.join(CLASSES_ROOT_DIR, "SS1-3.csv")

    if os.path.exists(js_csv):
        generate_word_report(js_csv, os.path.join(CLASSES_ROOT_DIR, "JS1-3_Report.docx"), "JUNIOR SECONDARY SCHOOL (JS1-3)")
    else:
        print("JS1-3.csv not found, cannot generate JS Word report.")

    if os.path.exists(ss_csv):
        generate_word_report(ss_csv, os.path.join(CLASSES_ROOT_DIR, "SS1-3_Report.docx"), "SENIOR SECONDARY SCHOOL (SS1-3)")
    else:
        print("SS1-3.csv not found, cannot generate SS Word report.")

def _get_overall_best_by_n(class_data_path, config):
    subject_lookup = _get_subject_lookup(config)

    # For each subject, keep list of students with the max score
    subject_best = {}  # subject_id -> {"max_score": int, "students": [{"name": str, "class": str}]}
    for file in Path(class_data_path).glob("*.json"):
        with open(file, "r") as f:
            data = json.load(f)

        cls = data.get("class") or {}
        class_label = f"{(cls.get('name') or '').strip()}{(cls.get('group') or '').strip()}"

        for item in data.get("top_students_in_subjects", []):
            subject_id = item.get("subject_id")
            name = (item.get("name") or "").strip()
            score = item.get("score", 0)

            if subject_id is None or not name:
                continue

            existing = subject_best.get(subject_id)
            if existing is None:
                subject_best[subject_id] = {"max_score": score, "students": [{"name": name, "class": class_label}]}
            else:
                if score > existing["max_score"]:
                    subject_best[subject_id] = {"max_score": score, "students": [{"name": name, "class": class_label}]}
                elif score == existing["max_score"]:
                    # Add this student as a co‑winner (avoid duplicates if same student appears twice)
                    if not any(s["name"] == name and s["class"] == class_label for s in existing["students"]):
                        subject_best[subject_id]["students"].append({"name": name, "class": class_label})

    # Group by student: count how many subjects they are best in
    student_subjects = {}  # key (name, class) -> list of subject names
    for subject_id, info in subject_best.items():
        subject_name = subject_lookup.get(subject_id, "Unknown")
        for student in info["students"]:
            key = (student["name"], student["class"])
            student_subjects.setdefault(key, []).append(subject_name)

    # Build overall_best_by_n: n -> {(name, class): [subjects]}
    overall_best_by_n = {}
    for key, subjects in student_subjects.items():
        n = len(subjects)
        overall_best_by_n.setdefault(n, {})[key] = subjects

    return overall_best_by_n


def _get_class_name_from_data_path(class_data_path):
    """Return the class name for the given data folder by reading the first JSON file."""
    for file in Path(class_data_path).glob("*.json"):
        with open(file, "r") as f:
            data = json.load(f)
        cls = data.get("class") or {}
        return (cls.get("name") or "").strip()
    return ""


def _generate_position_report_for_class(
    data,
    config=None,
    positions=None,
    avg_min=None,
    avg_max=None,
    best_scope=None,
    best_n=None,
    overall_best_by_n=None,
    filename=None,
):
    if config is None:
        config = load_config()
    if positions is None:
        positions = []
    """Generic engine for generating a per-class position/average based report.

    - positions: list of position keys (e.g. ['second', 'third']).
    - avg_min/avg_max: inclusive lower bound and exclusive upper bound for average.
    - best_scope: None|'local'|'overall'.
    - best_n: if best_scope is set, this is the exact number of subjects.
    - overall_best_by_n: precomputed map from n -> {(name, class): [subjects]}.
    - filename: output filename (required).
    """

    if filename is None:
        return

    class_name, class_group, class_label = _get_class_info(data)
    top_students = data.get("top_students", {}) or {}

    subject_lookup = _get_subject_lookup(config)

    # If we need local "best in N subjects" info, compute it once per class.
    student_subjects = None
    if best_scope == "local":
        student_subjects = _count_best_subjects_per_student(
            data.get("top_students_in_subjects", [])
        )

    # If we need overall-best info, make sure it's been computed.
    qualified_overall = None
    if best_scope == "overall":
        if overall_best_by_n is None:
            return
        qualified_overall = overall_best_by_n.get(best_n, {})

    result = []

    for position in positions:
        student = top_students.get(position)
        if not student:
            continue

        name = (student.get("name") or "").strip()
        average = student.get("average", 0)

        if avg_min is not None and average < avg_min:
            continue
        if avg_max is not None and average >= avg_max:
            continue

        subject_names = None

        if best_scope == "local":
            subjects = student_subjects.get(name, set())
            if best_n is not None and len(subjects) != best_n:
                continue
            subject_names = [
                subject_lookup.get(sid, "Unknown")
                for sid in subjects
            ]

        elif best_scope == "overall":
            key = (name, class_label)
            subjects = qualified_overall.get(key)
            if not subjects:
                continue
            subject_names = subjects

        row = {
            "name": name,
            "average": average,
            "position": position,
            "class": class_label,
        }

        if subject_names is not None:
            row["subject_name"] = ", ".join(subject_names)

        result.append(row)

    if not result:
        return

    field_names = ["name", "average", "position"]
    if any("subject_name" in r for r in result):
        field_names.insert(3, "subject_name")
    field_names.append("class")

    _write_class_csv(class_name, class_group, filename, result, field_names)


def calculate_students_best_in_n_subjects(data, config, n):
    class_name, class_group, _ = _get_class_info(data)

    subject_lookup = _get_subject_lookup(config)
    student_subjects = _count_best_subjects_per_student(
        data.get("top_students_in_subjects", [])
    )

    result = []
    for name, subjects in student_subjects.items():
        if len(subjects) != n:
            continue

        subject_names = [
            subject_lookup.get(sid, "Unknown")
            for sid in subjects
        ]

        result.append({
            "name": name,
            "subject_name": ", ".join(subject_names)
        })

    if not result:
        return

    _write_class_csv(
        class_name,
        class_group,
        f"best_in_{n}_subjects.csv",
        result,
        ["name", "subject_name"],
    )

def extract_and_save_most_improved_students(data):
    class_name = data.get("class")["name"]
    class_group = data.get("class")["group"]

    improved_students = data.get("most_improved_students", [])

    result = []

    for student in improved_students:
        result.append({
            "name": student.get("name", "").strip(),
            "improvement": student.get("improvement")
        })

    if not result:
        return
    
    save_to_csv(
        data=result,
        file_destination=os.path.join(
            CLASSES_ROOT_DIR,
            f"{class_name}/results/{class_name}_{class_group}/most_improved_students.csv"
        ),
        field_names=["name", "improvement"]
    )

# BEST IN N SUBJECTS, 1ST AMONG THE BEST (overall best across class level)
def calculate_overall_best_in_n_subjects(class_data_path, config, n):
    class_name = _get_class_name_from_data_path(class_data_path)
    if not class_name:
        return

    overall_best_by_n = _get_overall_best_by_n(class_data_path, config)
    entries = overall_best_by_n.get(n)
    if not entries:
        return

    result = [
        {
            "name": name,
            "subject_name": ", ".join(subjects),
            "class": class_label,
        }
        for (name, class_label), subjects in entries.items()
    ]

    _write_overall_csv(
        class_name,
        f"best_overall_in_{n}_subjects.csv",
        result,
        ["name", "subject_name", "class"],
    )

# 2ND OR 3RD POSITION BELOW AVERAGE 80 ONLY
def calculate_2nd_and_3rd_position_below_average_80(data):
    _generate_position_report_for_class(
        data=data,
        positions=["second", "third"],
        avg_min=70,
        avg_max=80,
        filename="2nd_or_3rd_position_below_average_80.csv",
    )

# 2ND OR 3RD POSITION BELOW AVERAGE 80, BEST IN ONE-FOUR SUBJECT
def calculate_2nd_and_3rd_position_below_average_80_best_in_n_subjects(data, config, n):
    _generate_position_report_for_class(
        data=data,
        config=config,
        positions=["second", "third"],
        avg_min=70,
        avg_max=80,
        best_scope="local",
        best_n=n,
        filename=f"2nd_or_3rd_position_below_average_80_best_in_{n}_subjects.csv",
    )

# 2ND OR 3RD POSITION BELOW AVERAGE 80, 1ST AMONG THE BEST, BEST IN ONE - SIX SUBJECTS
def calculate_2nd_3rd_below_80_and_overall_best_in_n_subjects(
    class_data_path,
    config,
    n,
    overall_best_by_n=None,
):
    if overall_best_by_n is None:
        overall_best_by_n = _get_overall_best_by_n(class_data_path, config)

    for file in Path(class_data_path).glob("*.json"):
        with open(file, "r") as f:
            data = json.load(f)

        _generate_position_report_for_class(
            data=data,
            config=config,
            positions=["second", "third"],
            avg_min=70,
            avg_max=80,
            best_scope="overall",
            best_n=n,
            overall_best_by_n=overall_best_by_n,
            filename=f"2nd_or_3rd_position_below_average_80_overall_best_in_{n}_subjects.csv",
        )

# 2ND OR 3RD POSITION ABOVE AVERAGE 80 ONLY
def calculate_2nd_and_3rd_position_above_average_80(data):
    _generate_position_report_for_class(
        data=data,
        positions=["second", "third"],
        avg_min=80,
        filename="2nd_or_3rd_position_above_average_80.csv",
    )

# 2ND OR 3RD POSITION ABOVE AVERAGE 80, BEST IN ONE-SIX SUBJECT
def calculate_2nd_and_3rd_position_above_average_80_best_in_n_subjects(data, config, n):
    _generate_position_report_for_class(
        data=data,
        config=config,
        positions=["second", "third"],
        avg_min=80,
        best_scope="local",
        best_n=n,
        filename=f"2nd_or_3rd_position_above_average_80_best_in_{n}_subjects.csv",
    )

# 2ND OR 3RD POSITION ABOVE AVERAGE 80, 1ST AMONG THE BEST, BEST IN ONE-FIVE SUBJECTS
def calculate_2nd_3rd_above_80_and_overall_best_in_n_subjects(
    class_data_path,
    config,
    n,
    overall_best_by_n=None,
):
    if overall_best_by_n is None:
        overall_best_by_n = _get_overall_best_by_n(class_data_path, config)

    for file in Path(class_data_path).glob("*.json"):
        with open(file, "r") as f:
            data = json.load(f)

        _generate_position_report_for_class(
            data=data,
            config=config,
            positions=["second", "third"],
            avg_min=80,
            best_scope="overall",
            best_n=n,
            overall_best_by_n=overall_best_by_n,
            filename=f"2nd_or_3rd_position_above_average_80_overall_best_in_{n}_subjects.csv",
        )

# 1ST POSITION BELOW AVERAGE 80 ONLY
def calculate_1st_position_below_average_80(data):
    _generate_position_report_for_class(
        data=data,
        positions=["first"],
        avg_min=70,
        avg_max=80,
        filename="1st_position_below_average_80.csv",
    )

# 1ST POSITION BELOW AVERAGE 80, BEST IN ONE-FIVE SUBJECT
def calculate_1st_position_below_average_80_best_in_n_subjects(data, config, n):
    _generate_position_report_for_class(
        data=data,
        config=config,
        positions=["first"],
        avg_min=70,
        avg_max=80,
        best_scope="local",
        best_n=n,
        filename=f"1st_position_below_average_80_best_in_{n}_subjects.csv",
    )

# 1ST POSITION BELOW AVERAGE 80, 1ST AMONG THE BEST, BEST IN ONE SUBJECT
def calculate_1st_below_80_and_overall_best_in_n_subjects(
    class_data_path,
    config,
    n,
    overall_best_by_n=None,
):
    if overall_best_by_n is None:
        overall_best_by_n = _get_overall_best_by_n(class_data_path, config)

    for file in Path(class_data_path).glob("*.json"):
        with open(file, "r") as f:
            data = json.load(f)

        _generate_position_report_for_class(
            data=data,
            config=config,
            positions=["first"],
            avg_min=70,
            avg_max=80,
            best_scope="overall",
            best_n=n,
            overall_best_by_n=overall_best_by_n,
            filename=f"1st_below_average_80_overall_best_in_{n}_subjects.csv",
        )

# 1ST POSITION ABOVE AVERAGE 80 ONLY
def calculate_1st_position_above_average_80(data):
    _generate_position_report_for_class(
        data=data,
        positions=["first"],
        avg_min=80,
        filename="1st_position_above_average_80.csv",
    )

# 1ST POSITION ABOVE AVERAGE 80, BEST IN ONE SUBJECT
def calculate_1st_position_above_average_80_best_in_n_subjects(data, config, n):
    _generate_position_report_for_class(
        data=data,
        config=config,
        positions=["first"],
        avg_min=80,
        best_scope="local",
        best_n=n,
        filename=f"1st_position_above_average_80_best_in_{n}_subjects.csv",
    )

# 1ST POSITION ABOVE AVERAGE 80, 1ST AMONG THE BEST, BEST IN ONE SUBJECT
def calculate_1st_above_80_and_overall_best_in_n_subjects(
    class_data_path,
    config,
    n,
    overall_best_by_n=None,
):
    if overall_best_by_n is None:
        overall_best_by_n = _get_overall_best_by_n(class_data_path, config)

    for file in Path(class_data_path).glob("*.json"):
        with open(file, "r") as f:
            data = json.load(f)

        _generate_position_report_for_class(
            data=data,
            config=config,
            positions=["first"],
            avg_min=80,
            best_scope="overall",
            best_n=n,
            overall_best_by_n=overall_best_by_n,
            filename=f"1st_above_average_80_overall_best_in_{n}_subjects.csv",
        )

def _read_csv_file(file_path):
    """Read a CSV and return list of dicts, or empty list if file missing."""
    if not os.path.exists(file_path):
        return []
    with open(file_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        return list(reader)

def generate_summary_csvs(class_data_path, config):
    """Generate JS1-3.csv and SS1-3.csv in the classes root directory."""
    # Build sections in the exact order of the template (only n‑ranges that are actually calculated)
    sections = []

    # 1. BEST IN N SUBJECTS ONLY (n=1..4)
    for n in range(1, 5):
        sections.append((
            f"BEST IN {n} SUBJECT{'S' if n>1 else ''} ONLY",
            f"best_in_{n}_subjects.csv",
            False,
            None
        ))

    # 2. THE MOST IMPROVED STUDENT
    sections.append(("THE MOST IMPROVED STUDENT", "most_improved_students.csv", False, None))

    # 3. 2ND/3RD POSITION BELOW AVERAGE 80 ONLY
    sections.append(("2ND OR 3RD POSITION BELOW AVERAGE 80 ONLY", "2nd_or_3rd_position_below_average_80.csv", False, None))

    # 4. 2ND/3RD BELOW 80, BEST IN N SUBJECTS (n=1..4)
    for n in range(1, 5):
        sections.append((
            f"2ND OR 3RD POSITION BELOW AVERAGE 80, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"2nd_or_3rd_position_below_average_80_best_in_{n}_subjects.csv",
            False,
            None
        ))

    # 5. BEST IN N SUBJECTS, 1ST AMONG THE BEST (n=1..6)
    for n in range(1, 7):
        sections.append((
            f"BEST IN {n} SUBJECT{'S' if n>1 else ''}, 1ST AMONG THE BEST",
            f"best_overall_in_{n}_subjects.csv",
            True,
            None
        ))

    # 6. 2ND/3RD BELOW 80, 1ST AMONG THE BEST, BEST IN N SUBJECTS (n=1..6)
    for n in range(1, 7):
        sections.append((
            f"2ND OR 3RD POSITION BELOW AVERAGE 80, 1ST AMONG THE BEST, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"2nd_or_3rd_position_below_average_80_overall_best_in_{n}_subjects.csv",
            True,
            None
        ))

    # 7. 2ND/3RD POSITION ABOVE AVERAGE 80 ONLY
    sections.append(("2ND OR 3RD POSITION ABOVE AVERAGE 80 ONLY", "2nd_or_3rd_position_above_average_80.csv", False, None))

    # 8. 2ND/3RD ABOVE 80, BEST IN N SUBJECTS (n=1..6)
    for n in range(1, 7):
        sections.append((
            f"2ND OR 3RD POSITION ABOVE AVERAGE 80, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"2nd_or_3rd_position_above_average_80_best_in_{n}_subjects.csv",
            False,
            None
        ))

    # 9. 2ND/3RD ABOVE 80, 1ST AMONG THE BEST, BEST IN N SUBJECTS (n=1..5)
    for n in range(1, 6):
        sections.append((
            f"2ND OR 3RD POSITION ABOVE AVERAGE 80, 1ST AMONG THE BEST, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"2nd_or_3rd_position_above_average_80_overall_best_in_{n}_subjects.csv",
            True,
            None
        ))

    # 10. 1ST POSITION BELOW AVERAGE 80 ONLY
    sections.append(("1ST POSITION BELOW AVERAGE 80 ONLY", "1st_position_below_average_80.csv", False, None))

    # 11. 1ST BELOW 80, BEST IN N SUBJECTS (n=1..5)
    for n in range(1, 6):
        sections.append((
            f"1ST POSITION BELOW AVERAGE 80, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"1st_position_below_average_80_best_in_{n}_subjects.csv",
            False,
            None
        ))

    # 12. 1ST BELOW 80, 1ST AMONG THE BEST, BEST IN N SUBJECTS (n=1..7)
    for n in range(1, 8):
        sections.append((
            f"1ST POSITION BELOW AVERAGE 80, 1ST AMONG THE BEST, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"1st_below_average_80_overall_best_in_{n}_subjects.csv",
            True,
            None
        ))

    # 13. 1ST POSITION ABOVE AVERAGE 80 ONLY
    sections.append(("1ST POSITION ABOVE AVERAGE 80 ONLY", "1st_position_above_average_80.csv", False, None))

    # 14. 1ST ABOVE 80, BEST IN N SUBJECTS (n=1..5)
    for n in range(1, 6):
        sections.append((
            f"1ST POSITION ABOVE AVERAGE 80, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"1st_position_above_average_80_best_in_{n}_subjects.csv",
            False,
            None
        ))

    # 15. 1ST ABOVE 80, 1ST AMONG THE BEST, BEST IN N SUBJECTS (n=1..8)
    for n in range(1, 9):
        sections.append((
            f"1ST POSITION ABOVE AVERAGE 80, 1ST AMONG THE BEST, BEST IN {n} SUBJECT{'S' if n>1 else ''}",
            f"1st_above_average_80_overall_best_in_{n}_subjects.csv",
            True,
            None
        ))

    # ------------------------------------------------------------------
    # Prepare containers for JS and SS
    js_rows = []
    ss_rows = []
    js_overall_candidates = []   # list of dicts {class_group, name, average}
    ss_overall_candidates = []
    js_handshake_rows = []
    ss_handshake_rows = []

    def add_rows_from_csv(section_name, class_group, csv_path, rows_list):
        data = _read_csv_file(csv_path)
        for row in data:
            out_row = {
                "Section": section_name,
                "Class": class_group,
                "Name": row.get("name", ""),
                "Subject(s)": row.get("subject_name", ""),
                "Average": row.get("average", ""),
                "Position": row.get("position", ""),
                "Improvement": row.get("improvement", ""),
            }
            rows_list.append(out_row)

    # Iterate over each class group JSON file
    for class_group_file in Path(class_data_path).glob("*.json"):
        with open(class_group_file, "r") as f:
            class_data = json.load(f)
        class_name = class_data["class"]["name"]
        class_group = f"{class_name}{class_data['class']['group']}"

        if class_name.startswith("JS"):
            target_rows = js_rows
            handshake_target = js_handshake_rows
            overall_target = js_overall_candidates
        elif class_name.startswith("SS"):
            target_rows = ss_rows
            handshake_target = ss_handshake_rows
            overall_target = ss_overall_candidates
        else:
            continue

        # 1. Add rows from all CSV sections
        for section_name, file_pattern, is_overall, _ in sections:
            if is_overall:
                overall_file = os.path.join(CLASSES_ROOT_DIR, class_name, "results", file_pattern)
                overall_data = _read_csv_file(overall_file)
                for row in overall_data:
                    if row.get("class") == class_group:
                        out_row = {
                            "Section": section_name,
                            "Class": class_group,
                            "Name": row.get("name", ""),
                            "Subject(s)": row.get("subject_name", ""),
                            "Average": "",
                            "Position": "",
                            "Improvement": "",
                        }
                        target_rows.append(out_row)
            else:
                csv_path = os.path.join(
                    CLASSES_ROOT_DIR, class_name, "results",
                    f"{class_name}_{class_data['class']['group']}",
                    file_pattern
                )
                add_rows_from_csv(section_name, class_group, csv_path, target_rows)

        # 2. Collect handshake rows (second and third positions)
        second = class_data.get("top_students", {}).get("second", {})
        third = class_data.get("top_students", {}).get("third", {})
        for pos, student in [("second", second), ("third", third)]:
            if student.get("name"):
                handshake_target.append({
                    "Section": "HAND SHAKE",
                    "Class": class_group,
                    "Name": student.get("name", ""),
                    "Subject(s)": "",
                    "Average": student.get("average", ""),
                    "Position": pos,
                    "Improvement": "",
                })

        # 3. Collect first‑position students for overall best
        first = class_data.get("top_students", {}).get("first", {})
        if first.get("name"):
            overall_target.append({
                "class_group": class_group,
                "name": first.get("name", ""),
                "average": first.get("average", 0),
            })

    # ------------------------------------------------------------------
    # Add OVERALL BEST STUDENTS (only the single top student per class level)
    for target_rows, candidates in [(js_rows, js_overall_candidates), (ss_rows, ss_overall_candidates)]:
        if candidates:
            best_student = max(candidates, key=lambda x: x.get("average", 0))
            target_rows.append({
                "Section": "OVERALL BEST STUDENTS",
                "Class": best_student["class_group"],
                "Name": best_student["name"],
                "Subject(s)": "",
                "Average": best_student["average"],
                "Position": "",
                "Improvement": "",
            })

    # Append handshake rows (they come after overall best in the template)
    js_rows.extend(js_handshake_rows)
    ss_rows.extend(ss_handshake_rows)

    # ------------------------------------------------------------------
    # Sort rows by the intended section order (based on the 'sections' list)
    section_order = {}
    for idx, (section_name, _, _, _) in enumerate(sections):
        section_order[section_name] = idx
    section_order["OVERALL BEST STUDENTS"] = len(sections)
    section_order["HAND SHAKE"] = len(sections) + 1

    for rows in (js_rows, ss_rows):
        rows.sort(key=lambda row: section_order.get(row.get("Section", ""), 9999))

    # ------------------------------------------------------------------
    # Write the two summary CSVs
    fieldnames = ["Section", "Class", "Name", "Subject(s)", "Average", "Position", "Improvement"]
    js_path = os.path.join(CLASSES_ROOT_DIR, "JS1-3.csv")
    ss_path = os.path.join(CLASSES_ROOT_DIR, "SS1-3.csv")

    for rows, path in [(js_rows, js_path), (ss_rows, ss_path)]:
        with open(path, "w", newline="", encoding="utf-8") as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)

    print(f"Summary CSVs generated: {js_path}, {ss_path}")

def execute_calculations(class_data_path):
    config = load_config()
    overall_best_by_n = _get_overall_best_by_n(class_data_path, config)
    for class_group in Path(class_data_path).iterdir():
        
        with open(class_group, 'r') as file:
            data = json.load(file)

            for n in range(1, 5):
                calculate_students_best_in_n_subjects(data, config, n)

            extract_and_save_most_improved_students(data)
            calculate_2nd_and_3rd_position_below_average_80(data)

            for n in range(1, 5):
                calculate_2nd_and_3rd_position_below_average_80_best_in_n_subjects(data, config, n)

            for n in range(1, 7):
                calculate_2nd_and_3rd_position_above_average_80_best_in_n_subjects(data, config, n)

            calculate_2nd_and_3rd_position_above_average_80(data)
            calculate_1st_position_below_average_80(data)

            for n in range(1, 6):
                calculate_1st_position_below_average_80_best_in_n_subjects(data, config, n)

            for n in range(1, 6):
                calculate_1st_position_above_average_80_best_in_n_subjects(data, config, n)

            calculate_1st_position_above_average_80(data)

            for n in range(1, 9):
                calculate_1st_above_80_and_overall_best_in_n_subjects(
                    class_data_path,
                    config,
                    n,
                    overall_best_by_n=overall_best_by_n,
                )

    for n in range(1, 7):
        calculate_2nd_3rd_below_80_and_overall_best_in_n_subjects(
            class_data_path,
            config,
            n,
            overall_best_by_n=overall_best_by_n,
        )

    for n in range(1, 6):
        calculate_2nd_3rd_above_80_and_overall_best_in_n_subjects(
            class_data_path,
            config,
            n,
            overall_best_by_n=overall_best_by_n,
        )

    for n in range(1, 7):
        calculate_overall_best_in_n_subjects(
            class_data_path=class_data_path,
            config=config,
            n=n
        ) 

    for n in range(1, 8):
        calculate_1st_below_80_and_overall_best_in_n_subjects(
            class_data_path,
            config,
            n,
            overall_best_by_n=overall_best_by_n,
        )

    generate_summary_csvs(class_data_path, config)
    generate_all_word_reports() 